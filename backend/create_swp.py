"""Build the Safe Work Procedure DOCX document."""
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11 if level == 1 else 10)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def _body(doc, text: str, indent: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def build_swp_docx(project_details: dict, swp_data: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ── Title block ────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SAFE WORK PROCEDURE")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(project_details.get("project_name", ""))
    r2.bold = True
    r2.font.size = Pt(12)

    rev = doc.add_paragraph()
    rev.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = rev.add_run(f"Rev 0    {project_details.get('assessment_date', datetime.today().strftime('%d %b %Y'))}")
    r3.font.size = Pt(10)

    doc.add_paragraph()

    # ── Header table ───────────────────────────────────────────────────────
    ht = doc.add_table(rows=2, cols=4)
    ht.style = "Table Grid"
    labels = [
        ["Prepared By:", project_details.get("ra_leader", ""),
         "Acknowledged By:", project_details.get("approved_by", "")],
        ["Company:", project_details.get("company", "GWS LIVINGART PTE LTD"),
         "Location:", project_details.get("location", swp_data.get("location", ""))],
    ]
    for r_idx, row_data in enumerate(labels):
        row = ht.rows[r_idx]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = (c_idx % 2 == 0)
            run.font.size = Pt(9)
            if c_idx % 2 == 0:
                _set_cell_bg(cell, "D9E1F2")

    doc.add_paragraph()

    # ── Purpose ────────────────────────────────────────────────────────────
    _heading(doc, "1.  Purpose")
    _body(doc, swp_data.get("purpose", "To provide a safe work procedure for the above-mentioned works."))

    # ── Roles and Responsibilities ─────────────────────────────────────────
    _heading(doc, "2.  Roles and Responsibilities")

    roles = [
        ("Manager / Director",
         "Overall in charge of project. Ensure work is carried out according to legal requirements and relevant Singapore Standards. Ensure Risk Assessment and Safe Work Procedure are in place. Provide necessary PPE, lifting equipment and access equipment required for the works."),
        ("WSH Supervisor / Coordinator",
         "Ensure all works are carried out safely. Ensure relevant Permit-to-Work is applied and approved before commencement. Conduct daily toolbox meeting and RA/SWP briefing. Carry out regular site safety inspections. Report any unsafe acts or unsafe conditions."),
        ("Site Supervisor / Foreman",
         "Apply Permit-to-Work. Supervise daily construction activities in accordance with approved plan. Ensure workers always wear required PPE. Comply with Main Contractor in-house rules. Report incidents/near-misses immediately."),
        ("Workers",
         "Follow all instructions from supervisors. Wear required PPE at all times. Stop work and report if unsafe condition is observed. Attend daily toolbox meeting. Do not carry out work without proper authorisation."),
    ]

    role_table = doc.add_table(rows=1 + len(roles), cols=2)
    role_table.style = "Table Grid"
    role_table.columns[0].width = Cm(5)
    role_table.columns[1].width = Cm(11)
    for i, hdr in enumerate(["Role", "Responsibilities"]):
        cell = role_table.rows[0].cells[i]
        cell.paragraphs[0].add_run(hdr).bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(cell, "1F3864")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, (role, resp) in enumerate(roles):
        row = role_table.rows[idx + 1]
        r0 = row.cells[0].paragraphs[0]
        r0.add_run(role).bold = True
        r0.runs[0].font.size = Pt(9)
        r1 = row.cells[1].paragraphs[0]
        r1.add_run(resp).font.size = Pt(9)
        if idx % 2 == 0:
            _set_cell_bg(row.cells[0], "EBF3FB")
            _set_cell_bg(row.cells[1], "EBF3FB")

    doc.add_paragraph()

    # ── Work Activities ────────────────────────────────────────────────────
    _heading(doc, "3.  Work Activities")

    activities = swp_data.get("activities", [])
    for a_idx, activity in enumerate(activities):
        # Activity heading
        act_p = doc.add_paragraph()
        act_p.paragraph_format.space_before = Pt(6)
        act_run = act_p.add_run(f"3.{a_idx + 1}  {activity.get('name', '')}")
        act_run.bold = True
        act_run.font.size = Pt(10)
        act_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        steps = activity.get("steps", [])
        for s_idx, step in enumerate(steps):
            step_p = doc.add_paragraph(style="List Number")
            step_p.paragraph_format.left_indent = Cm(1)
            step_p.paragraph_format.space_before = Pt(1)
            step_p.paragraph_format.space_after = Pt(1)
            run = step_p.add_run(step)
            run.font.size = Pt(9)

    doc.add_paragraph()

    # ── PPE Requirements ───────────────────────────────────────────────────
    _heading(doc, "4.  Mandatory PPE Requirements")

    ppe_items = [
        "Safety Helmet (compliant with SS 98)",
        "Safety Boots / Shoes (steel-toed)",
        "Safety Vest / High-Visibility Vest",
        "Work Gloves",
        "Safety Spectacles / Goggles (where required)",
        "Full Body Harness with double lanyard (for Work at Height activities)",
        "Dust Mask / Respirator (where required)",
        "Ear Defenders (where noise exposure exceeds 85 dB)",
    ]
    for item in ppe_items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item).font.size = Pt(9)

    doc.add_paragraph()

    # ── Emergency Procedures ───────────────────────────────────────────────
    _heading(doc, "5.  Emergency Procedures")
    emergencies = [
        "In case of injury: Render first aid immediately. Call 995 for ambulance if required. Inform Site Supervisor and Safety Officer immediately.",
        "In case of fire: Activate fire alarm. Call 995. Evacuate all personnel to assembly point. Do not re-enter until declared safe.",
        "In case of fall from height: Do not move the injured person. Call 995. Secure the area. Inform Site Supervisor immediately.",
        "Emergency Assembly Point: As directed by Main Contractor site safety signage.",
        "SGSecure: Be vigilant for suspicious items or activities. Report to police (999) if suspicious activity is observed.",
    ]
    for item in emergencies:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item).font.size = Pt(9)

    doc.add_paragraph()

    # ── Acknowledgement table ──────────────────────────────────────────────
    _heading(doc, "6.  Worker Acknowledgement")
    ack_p = doc.add_paragraph()
    ack_p.add_run("I have read and understood the above Safe Work Procedure and agree to comply.").font.size = Pt(9)

    ack_table = doc.add_table(rows=6, cols=4)
    ack_table.style = "Table Grid"
    for i, hdr in enumerate(["No.", "Name", "Signature", "Date"]):
        cell = ack_table.rows[0].cells[i]
        cell.paragraphs[0].add_run(hdr).bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(cell, "D9E1F2")
    for r_idx in range(1, 6):
        row = ack_table.rows[r_idx]
        row.cells[0].paragraphs[0].add_run(str(r_idx)).font.size = Pt(9)
        for ci in range(1, 4):
            row.cells[ci].paragraphs[0].add_run(" ").font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
