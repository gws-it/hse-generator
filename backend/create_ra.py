"""Build the Risk Assessment DOCX document."""
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell(cell, text: str, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))


def _rpn_color(rpn: int) -> str:
    if rpn <= 4:
        return "92D050"   # green
    elif rpn <= 9:
        return "FFFF00"   # yellow
    elif rpn <= 16:
        return "FFC000"   # orange
    else:
        return "FF0000"   # red


def build_ra_docx(project_details: dict, ra_data: dict) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width = Cm(42)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.orientation = 1  # landscape

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("EHS RISK ASSESSMENT")
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"{project_details.get('project_name', '')} – {project_details.get('company', 'GWS LIVINGART PTE LTD')}")
    r.font.size = Pt(10)
    r.bold = True

    doc.add_paragraph()

    # ── Header info table ──────────────────────────────────────────────────
    h = doc.add_table(rows=4, cols=6)
    h.style = "Table Grid"
    hdr_data = [
        ["Company:", project_details.get("company", "GWS LIVINGART PTE LTD"),
         "RA Leader:", project_details.get("ra_leader", ""),
         "Reference No.", project_details.get("reference_no", "")],
        ["Process:", project_details.get("project_name", ""),
         "RA Members 1:", (project_details.get("ra_members") or [""])[0],
         "Original Assessment Date:", project_details.get("assessment_date", datetime.today().strftime("%d %b %Y"))],
        ["Activity Location:", project_details.get("location", ""),
         "RA Members 2:", (project_details.get("ra_members") or ["", ""])[1] if len(project_details.get("ra_members") or []) > 1 else "",
         "Last Review Date:", project_details.get("assessment_date", datetime.today().strftime("%d %b %Y"))],
        ["Approved By:", project_details.get("approved_by", ""),
         "RA Members 3:", (project_details.get("ra_members") or ["", "", ""])[2] if len(project_details.get("ra_members") or []) > 2 else "",
         "Next Review Date:", ""],
    ]
    for r_idx, row_data in enumerate(hdr_data):
        row = h.rows[r_idx]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            bold = c_idx % 2 == 0
            _cell(cell, text, bold=bold, size=8)

    doc.add_paragraph()

    # ── Activity list table ────────────────────────────────────────────────
    activities = ra_data.get("activities", [])

    al = doc.add_table(rows=1 + len(activities), cols=4)
    al.style = "Table Grid"
    headers = ["#", "S/N", "R/NR", "Sub-Activity"]
    widths = [Cm(1), Cm(1.5), Cm(1.5), Cm(10)]
    for i, h_text in enumerate(headers):
        cell = al.rows[0].cells[i]
        _cell(cell, h_text, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(cell, "D9E1F2")
        al.columns[i].width = widths[i]

    for idx, act in enumerate(activities):
        row = al.rows[idx + 1]
        _cell(row.cells[0], idx + 1, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[1], act.get("sn", ""), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[2], "R", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[3], act.get("sub_activity", ""), size=8)

    doc.add_paragraph()

    # ── RA Table ───────────────────────────────────────────────────────────
    col_headers = [
        "S/N", "Sub-Activity", "Hazard / Aspect",
        "Possible Injury / Ill-health\nDamage / Env Impact",
        "Existing Control Measures", "S", "L", "RPN",
        "Additional Control Measures", "S", "L", "RPN",
        "Implementation Person", "Due Date", "Remarks"
    ]
    col_widths = [
        Cm(1.2), Cm(3.5), Cm(3.5), Cm(3.5),
        Cm(5.5), Cm(0.8), Cm(0.8), Cm(1.0),
        Cm(4.5), Cm(0.8), Cm(0.8), Cm(1.0),
        Cm(2.5), Cm(2.0), Cm(2.0)
    ]

    ra_table = doc.add_table(rows=1 + len(activities), cols=len(col_headers))
    ra_table.style = "Table Grid"

    # Header row
    for i, h_text in enumerate(col_headers):
        cell = ra_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(7)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, "1F3864")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        ra_table.columns[i].width = col_widths[i]

    # Data rows
    for idx, act in enumerate(activities):
        row = ra_table.rows[idx + 1]
        ec = act.get("existing_controls", {})
        ac = act.get("additional_controls", {})

        def fmt_controls(c: dict) -> str:
            parts = []
            for key in ["elimination", "substitution", "engineering", "administrative", "ppe"]:
                val = c.get(key, "NA")
                if val and val.upper() != "NA":
                    parts.append(f"{key.capitalize()}:\n{val}")
            return "\n\n".join(parts) if parts else "NA"

        i_rpn = act.get("initial_rpn", act.get("initial_s", 1) * act.get("initial_l", 1))
        r_rpn = act.get("residual_rpn", act.get("residual_s", 1) * act.get("residual_l", 1))

        cells_data = [
            (act.get("sn", ""), False, WD_ALIGN_PARAGRAPH.CENTER),
            (act.get("sub_activity", ""), True, WD_ALIGN_PARAGRAPH.LEFT),
            (act.get("hazard", ""), False, WD_ALIGN_PARAGRAPH.LEFT),
            (act.get("possible_injury", ""), False, WD_ALIGN_PARAGRAPH.LEFT),
            (fmt_controls(ec), False, WD_ALIGN_PARAGRAPH.LEFT),
            (act.get("initial_s", ""), False, WD_ALIGN_PARAGRAPH.CENTER),
            (act.get("initial_l", ""), False, WD_ALIGN_PARAGRAPH.CENTER),
            (i_rpn, True, WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_controls(ac), False, WD_ALIGN_PARAGRAPH.LEFT),
            (act.get("residual_s", ""), False, WD_ALIGN_PARAGRAPH.CENTER),
            (act.get("residual_l", ""), False, WD_ALIGN_PARAGRAPH.CENTER),
            (r_rpn, True, WD_ALIGN_PARAGRAPH.CENTER),
            (act.get("implementation_person", "Site Supervisor"), False, WD_ALIGN_PARAGRAPH.CENTER),
            (act.get("due_date", "On-going"), False, WD_ALIGN_PARAGRAPH.CENTER),
            (act.get("remarks", ""), False, WD_ALIGN_PARAGRAPH.LEFT),
        ]

        for c_idx, (text, bold, align) in enumerate(cells_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.alignment = align
            run = p.add_run(str(text))
            run.bold = bold
            run.font.size = Pt(7)

        # Colour-code RPN cells
        _set_cell_bg(row.cells[7], _rpn_color(int(i_rpn)))
        _set_cell_bg(row.cells[11], _rpn_color(int(r_rpn)))

        # Alternate row shading
        if idx % 2 == 1:
            for ci in [0, 1, 2, 3, 4, 8, 12, 13, 14]:
                _set_cell_bg(row.cells[ci], "EBF3FB")

    # ── Signature block ────────────────────────────────────────────────────
    doc.add_paragraph()
    sig = doc.add_table(rows=2, cols=3)
    sig.style = "Table Grid"
    sig_headers = ["Prepared By", "Reviewed By", "Approved By"]
    sig_names = [
        project_details.get("ra_leader", ""),
        "",
        project_details.get("approved_by", ""),
    ]
    for i, h_text in enumerate(sig_headers):
        _cell(sig.rows[0].cells[i], h_text, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(sig.rows[0].cells[i], "D9E1F2")
    for i, name in enumerate(sig_names):
        cell = sig.rows[1].cells[i]
        cell.paragraphs[0].add_run("\n\n\n")
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(name).font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
